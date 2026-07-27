"""
Markdown to DOCX Converter for Project Reports.
Converts Markdown files (.md) in reports/ to Microsoft Word documents (.docx) with professional styling,
tables, headers, bullet points, callout blocks, and code blocks.
"""

import os
import re
import sys

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install --break-system-packages python-docx")
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def markdown_to_docx(md_filepath, docx_filepath):
    print(f"Converting {md_filepath} -> {docx_filepath}...")
    doc = Document()

    # Base Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set normal style font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []

    def flush_table(rows):
        if not rows:
            return
        headers = rows[0]
        data_rows = [r for r in rows[1:] if not all(c.strip('-: ') == '' for c in r)]

        if not headers:
            return

        tbl = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        # Format Header Row
        hdr_cells = tbl.rows[0].cells
        for i, text in enumerate(headers):
            hdr_cells[i].text = text.strip('*_#`')
            set_cell_background(hdr_cells[i], "1E293B") # Dark navy header
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)
                run.font.size = Pt(10)

        # Format Data Rows
        for r_idx, row_data in enumerate(data_rows):
            row_cells = tbl.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < len(row_cells):
                    row_cells[c_idx].text = cell_text.strip()
                    set_cell_background(row_cells[c_idx], bg_color)
                    p = row_cells[c_idx].paragraphs[0]
                    for run in p.runs:
                        run.font.size = Pt(9.5)

        doc.add_paragraph() # Spacing after table

    for line in lines:
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                code_text = "".join(code_block_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                code_block_lines = []
            else:
                if in_table:
                    flush_table(table_rows)
                    in_table = False
                    table_rows = []
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Handle Tables
        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            cols = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cols)
            continue
        elif in_table:
            flush_table(table_rows)
            in_table = False
            table_rows = []

        if not stripped:
            continue

        # Headings
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            run = h.add_run(stripped[2:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[3:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        elif stripped.startswith("#### "):
            h = doc.add_heading(level=4)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(stripped[5:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

        # Blockquotes / Callout Alerts
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            text = stripped[2:].strip()
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        # Bullet Lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Basic bold formatting check
            clean_text = stripped[2:].strip()
            parts = re.split(r'(\*\*.*?\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Paragraph Text
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            clean_text = stripped
            parts = re.split(r'(\*\*.*?\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    if in_table:
        flush_table(table_rows)

    doc.save(docx_filepath)
    print(f"Successfully created: {docx_filepath}")

if __name__ == "__main__":
    reports_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")
    md_files = [
        "financial_inclusion_forecasting_report.md",
        "blog_post_medium.md",
        "interim_progress_report.md",
        "assignment_rubric_evaluation.md",
        "progress_report_rubric_evaluation.md"
    ]

    for md_file in md_files:
        md_path = os.path.join(reports_dir, md_file)
        if os.path.exists(md_path):
            docx_path = os.path.join(reports_dir, md_file.replace(".md", ".docx"))
            markdown_to_docx(md_path, docx_path)

